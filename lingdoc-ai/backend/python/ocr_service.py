#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PaddleOCR 识别服务脚本

功能：
1. 接收文件路径参数，支持 PDF、Word(docx/doc)、图片(jpg/png/bmp)
2. Word/PDF 自动转图片后识别
3. 输出结构化 JSON 结果（含文本、置信度、位置信息）
4. 纯 CPU 运行，支持中文识别

用法（命令行）：
    python3 ocr_service.py --file /path/to/doc.pdf --dpi 150 --threshold 0.5 --format json

用法（HTTP服务模式，预留）：
    python3 ocr_service.py --server --port 5000
"""

import os
import sys
import time
import json
import argparse
import tempfile
import warnings

# 抑制不必要的警告
warnings.filterwarnings("ignore")

# ============ 依赖检查与安装提示 ============
try:
    from paddleocr import PaddleOCR
except ImportError:
    print(json.dumps({
        "error": "PaddleOCR 未安装，请先执行: pip install paddleocr==2.8.1"
    }, ensure_ascii=False))
    sys.exit(1)

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image
except ImportError:
    Image = None


# ============ 全局配置 ============
OCR_ENGINE = None  # 延迟初始化


def get_ocr_engine(use_gpu=False, lang="ch", show_log=False):
    """获取或初始化 PaddleOCR 引擎（单例模式）"""
    global OCR_ENGINE
    if OCR_ENGINE is None:
        OCR_ENGINE = PaddleOCR(
            use_angle_cls=True,
            lang=lang,
            show_log=show_log,
            use_gpu=use_gpu,
            det_db_box_thresh=0.5,
            rec_thresh=0.5
        )
    return OCR_ENGINE


def convert_pdf_to_images(pdf_path, dpi=150):
    """PDF 转图片列表"""
    if fitz is None:
        raise RuntimeError("PyMuPDF 未安装，无法处理 PDF: pip install pymupdf")

    images = []
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=dpi)
        img_path = os.path.join(tempfile.gettempdir(), f"ocr_page_{page_num}_{int(time.time())}.png")
        pix.save(img_path)
        images.append(img_path)
    doc.close()
    return images


def convert_docx_to_images(docx_path, dpi=150):
    """Word 文档转图片（通过 PDF 中转）"""
    if fitz is None:
        raise RuntimeError("PyMuPDF 未安装，无法处理 Word 文档")

    # 使用 LibreOffice 或 unoconv 将 Word 转 PDF
    pdf_path = os.path.join(tempfile.gettempdir(), f"ocr_tmp_{int(time.time())}.pdf")

    # 优先尝试 libreoffice
    cmd_libre = f"libreoffice --headless --convert-to pdf --outdir {os.path.dirname(pdf_path)} {docx_path}"
    ret = os.system(cmd_libre)

    if ret != 0:
        # 备选：直接读取文本（无损但无布局）
        if Document is not None:
            doc = Document(docx_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return ["TEXT_MODE"], text
        raise RuntimeError("无法转换 Word 文档，请安装 LibreOffice 或 python-docx")

    # 找到生成的 PDF
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    generated_pdf = os.path.join(os.path.dirname(pdf_path), base_name + ".pdf")
    if not os.path.exists(generated_pdf):
        raise RuntimeError("LibreOffice 转换失败，未生成 PDF 文件")

    images = convert_pdf_to_images(generated_pdf, dpi)
    os.remove(generated_pdf)
    return images


def recognize_image(ocr, image_path, page_num=1):
    """单张图片 OCR 识别"""
    result = ocr.ocr(image_path, cls=True)

    lines = []
    if result and result[0]:
        for line in result[0]:
            bbox = line[0]  # 边界框坐标
            text_info = line[1]  # (text, confidence)
            lines.append({
                "text": text_info[0],
                "confidence": float(text_info[1]),
                "bbox": [float(c) for point in bbox for c in point]  # 展平为 [x1,y1,x2,y2,...]
            })

    return {
        "page_num": page_num,
        "line_count": len(lines),
        "lines": lines
    }


def recognize_text(text, page_num=1):
    """文本模式识别（Word 直接提取文本，无 OCR）"""
    lines = [{"text": line, "confidence": 1.0, "bbox": []} for line in text.split("\n") if line.strip()]
    return {
        "page_num": page_num,
        "line_count": len(lines),
        "lines": lines,
        "mode": "text_extraction"
    }


def process_file(file_path, dpi=150, confidence_threshold=0.5):
    """处理文件并返回结构化结果"""
    start_time = time.time()
    file_ext = os.path.splitext(file_path)[1].lower()
    ocr = get_ocr_engine()

    pages = []
    total_chars = 0
    total_lines = 0
    process_mode = "ocr"

    try:
        if file_ext in [".pdf"]:
            images = convert_pdf_to_images(file_path, dpi)
            for i, img_path in enumerate(images):
                page_result = recognize_image(ocr, img_path, page_num=i + 1)
                pages.append(page_result)
                total_lines += page_result["line_count"]
                for line in page_result["lines"]:
                    total_chars += len(line["text"])
                os.remove(img_path)  # 清理临时文件

        elif file_ext in [".docx", ".doc"]:
            try:
                images = convert_docx_to_images(file_path, dpi)
                if isinstance(images, tuple) and images[0] == ["TEXT_MODE"]:
                    # 文本模式
                    pages.append(recognize_text(images[1], page_num=1))
                    total_lines = pages[0]["line_count"]
                    total_chars = sum(len(l["text"]) for l in pages[0]["lines"])
                    process_mode = "text_extraction"
                else:
                    for i, img_path in enumerate(images):
                        page_result = recognize_image(ocr, img_path, page_num=i + 1)
                        pages.append(page_result)
                        total_lines += page_result["line_count"]
                        for line in page_result["lines"]:
                            total_chars += len(line["text"])
                        os.remove(img_path)
            except Exception as e:
                # Word 转换失败时回退到文本提取
                if Document is not None:
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    pages.append(recognize_text(text, page_num=1))
                    total_lines = pages[0]["line_count"]
                    total_chars = sum(len(l["text"]) for l in pages[0]["lines"])
                    process_mode = "text_extraction_fallback"
                else:
                    raise e

        elif file_ext in [".jpg", ".jpeg", ".png", ".bmp"]:
            page_result = recognize_image(ocr, file_path, page_num=1)
            pages.append(page_result)
            total_lines = page_result["line_count"]
            total_chars = sum(len(l["text"]) for l in page_result["lines"])

        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")

    except Exception as e:
        return {
            "error": str(e),
            "file_path": file_path,
            "process_time_ms": int((time.time() - start_time) * 1000)
        }

    process_time = int((time.time() - start_time) * 1000)
    avg_page_time = process_time // len(pages) if pages else 0

    # 过滤低置信度结果
    for page in pages:
        if "lines" in page:
            page["lines"] = [
                line for line in page["lines"]
                if line.get("confidence", 0) >= confidence_threshold
            ]
            page["line_count"] = len(page["lines"])

    return {
        "file_path": file_path,
        "file_type": file_ext.replace(".", ""),
        "page_count": len(pages),
        "total_lines": total_lines,
        "total_chars": total_chars,
        "process_time_ms": process_time,
        "avg_page_time_ms": avg_page_time,
        "process_mode": process_mode,
        "engine": "PaddleOCR",
        "engine_version": "2.8.1",
        "pages": pages
    }


def main():
    parser = argparse.ArgumentParser(description="PaddleOCR 文档识别服务")
    parser.add_argument("--file", "-f", type=str, help="待识别文件路径")
    parser.add_argument("--dpi", type=int, default=150, help="PDF转图片DPI（默认150）")
    parser.add_argument("--threshold", "-t", type=float, default=0.5, help="置信度阈值（默认0.5）")
    parser.add_argument("--format", choices=["json", "text"], default="json", help="输出格式")
    parser.add_argument("--server", action="store_true", help="启动HTTP服务模式")
    parser.add_argument("--port", type=int, default=5000, help="HTTP服务端口")
    parser.add_argument("--gpu", action="store_true", help="使用GPU（默认CPU）")

    args = parser.parse_args()

    if args.server:
        # HTTP 服务模式（预留）
        try:
            from flask import Flask, request, jsonify
            app = Flask(__name__)

            @app.route("/ocr", methods=["POST"])
            def ocr_endpoint():
                data = request.get_json()
                file_path = data.get("file_path")
                if not file_path or not os.path.exists(file_path):
                    return jsonify({"error": "文件不存在"}), 400
                dpi = data.get("dpi", 150)
                threshold = data.get("threshold", 0.5)
                result = process_file(file_path, dpi, threshold)
                return jsonify(result)

            print(f"[*] PaddleOCR HTTP 服务启动于 http://0.0.0.0:{args.port}")
            app.run(host="0.0.0.0", port=args.port, threaded=True)
        except ImportError:
            print(json.dumps({
                "error": "HTTP 模式需要 Flask，请先安装: pip install flask"
            }, ensure_ascii=False))
            sys.exit(1)
        return

    if not args.file:
        parser.print_help()
        sys.exit(1)

    if not os.path.exists(args.file):
        print(json.dumps({"error": f"文件不存在: {args.file}"}, ensure_ascii=False))
        sys.exit(1)

    # 命令行模式：直接输出 JSON
    result = process_file(args.file, args.dpi, args.threshold)

    if args.format == "json":
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        # 纯文本输出
        for page in result.get("pages", []):
            print(f"\n===== 第 {page['page_num']} 页 =====")
            for line in page.get("lines", []):
                print(f"[{line['confidence']:.2f}] {line['text']}")


if __name__ == "__main__":
    main()
