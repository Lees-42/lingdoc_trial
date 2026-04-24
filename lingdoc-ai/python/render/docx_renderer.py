#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 表格智能填表渲染器 v3
修复：合并单元格重复填充问题（用底层 _tc id 去重）
"""

import sys
import json
from docx import Document


def clean_text(text: str) -> str:
    if not text:
        return ""
    return text.strip().replace("：", "").replace(":", "").replace(" ", "").replace("\n", "").replace("\t", "").replace("[", "").replace("]", "")


def is_placeholder(text: str) -> bool:
    if not text:
        return True
    t = text.strip()
    if t.startswith("[") and t.endswith("]"):
        return True
    if "请" in t and ("填写" in t or "陈述" in t or "输入" in t or "不少于" in t):
        return True
    if t.replace("_", "").replace("—", "").replace("-", "").replace(" ", "") == "":
        return True
    return False


def _find_cell_in_table(table, field_name: str, filled_cell_ids: set):
    """
    查找字段名对应的目标值单元格
    filled_cell_ids: 已填充单元格的底层 XML _tc id 集合（解决合并单元格问题）
    """
    clean_name = clean_text(field_name)
    if not clean_name:
        return None, None

    def _tc_id(cell):
        return id(cell._tc)

    def _is_available(cell):
        return is_placeholder(cell.text) and _tc_id(cell) not in filled_cell_ids

    # 策略1: 直接匹配字段名，取右侧相邻的占位符单元格
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell_text = clean_text(cell.text)
            if cell_text == clean_name:
                if c_idx + 1 < len(row.cells):
                    right = row.cells[c_idx + 1]
                    if _is_available(right):
                        return right, "right_placeholder"
                if r_idx + 1 < len(table.rows):
                    down = table.rows[r_idx + 1].cells[c_idx]
                    if _is_available(down):
                        return down, "down_placeholder"

    # 策略2: 模糊匹配
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell_text = clean_text(cell.text)
            if clean_name in cell_text or cell_text in clean_name:
                if c_idx + 1 < len(row.cells):
                    right = row.cells[c_idx + 1]
                    if _is_available(right):
                        return right, "fuzzy_right"
                if r_idx + 1 < len(table.rows):
                    down = table.rows[r_idx + 1].cells[c_idx]
                    if _is_available(down):
                        return down, "fuzzy_down"

    # 策略3: 直接匹配 [字段名] 占位符
    placeholder = f"[{field_name}]"
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if placeholder in cell.text or f"[{clean_text(field_name)}]" in cell.text:
                if _tc_id(cell) not in filled_cell_ids:
                    return cell, "direct_placeholder"

    # 策略4: 匹配清洗后的 [字段名]
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if f"[{clean_text(field_name)}]" == clean_text(cell.text):
                if _tc_id(cell) not in filled_cell_ids:
                    return cell, "cleaned_placeholder"

    return None, None


def fill_word_document(input_path: str, output_path: str, filled_values: dict):
    doc = Document(input_path)
    filled_count = 0
    unmatched = []
    matched = []
    filled_cell_ids = set()  # 底层 _tc id 去重

    for table_idx, table in enumerate(doc.tables):
        for field_name, field_value in filled_values.items():
            target_cell, match_type = _find_cell_in_table(table, field_name, filled_cell_ids)
            if target_cell is not None:
                _set_cell_text(target_cell, str(field_value) if field_value is not None else "")
                filled_cell_ids.add(id(target_cell._tc))
                filled_count += 1
                matched.append({
                    "fieldName": field_name,
                    "value": field_value,
                    "table": table_idx,
                    "matchType": match_type
                })
            else:
                if field_name not in [u["fieldName"] for u in unmatched]:
                    unmatched.append({"fieldName": field_name, "reason": "未找到对应占位符"})

    # 段落中的 {{xxx}} 占位符
    for para in doc.paragraphs:
        for field_name, field_value in filled_values.items():
            placeholder = f"{{{{{field_name}}}}}"
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, str(field_value) if field_value else "")
                filled_count += 1
                matched.append({"fieldName": field_name, "value": field_value, "location": "paragraph"})

    doc.save(output_path)
    return {
        "success": True,
        "filledCount": filled_count,
        "matchedFields": matched,
        "unmatchedFields": unmatched,
        "outputPath": output_path
    }


def _set_cell_text(cell, text):
    """设置单元格文本，尽量保留格式"""
    first_para = cell.paragraphs[0] if cell.paragraphs else None
    if first_para:
        if first_para.runs:
            first_para.runs[0].text = text
            for run in first_para.runs[1:]:
                run.text = ""
        else:
            first_para.add_run(text)
    else:
        cell.text = text


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python docx_renderer.py <input.docx> <output.docx> '{\"姓名\":\"张三\"}'")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]
    filled_values = json.loads(sys.argv[3])

    result = fill_word_document(input_path, output_path, filled_values)
    print(json.dumps(result, ensure_ascii=False, indent=2))
