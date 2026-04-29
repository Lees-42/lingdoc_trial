# -*- coding: utf-8 -*-
"""
=============================================================================
Form Service - 智能填表业务逻辑
=============================================================================
替代 Dify 工作流：form-extract + form-generate + render

职责:
  1. 从参考文档提取结构化信息 (form-extract)
  2. 将提取信息匹配到表格字段 (form-generate)
  3. 调用渲染器生成最终文件 (render)

模型: qwen3-max (DashScope)
=============================================================================
"""

import json
import os
import re
import time
from typing import Dict, List, Optional, Tuple

from app.config import config
from app.utils.logger import logger
from app.services.llm_client import LLMClient
from app.prompts.form_prompts import FORM_EXTRACT_PROMPT, FORM_GENERATE_PROMPT


class FormService:
    """智能填表服务"""
    
    def __init__(self):
        self.llm = LLMClient()
    
    # =====================================================================
    # Step 1: 从参考文档提取信息 (替代 Dify form-extract)
    # =====================================================================
    
    async def extract_from_document(self, file_path: str, task_id: str) -> Dict:
        """
        从单个参考文档提取结构化信息
        
        Args:
            file_path: 参考文档绝对路径 (.docx / .pdf / .txt)
            task_id: 任务ID，用于日志追踪
            
        Returns:
            {"success": bool, "data": dict, "error": str, "task_id": str}
        """
        logger.info(f"[EXTRACT:{task_id}] 开始提取: {file_path}")
        start = time.time()
        
        try:
            # 1. 提取纯文本（非 OCR，直接读取）
            text = self._extract_text(file_path)
            if not text or len(text.strip()) < 5:
                return {
                    "success": False,
                    "data": {},
                    "error": "文档内容为空或无法读取",
                    "task_id": task_id
                }
            
            logger.info(f"[EXTRACT:{task_id}] 文本提取完成 | 长度={len(text)}")
            
            # 2. 调用 LLM 提取结构化信息
            prompt = FORM_EXTRACT_PROMPT.format(document_text=text[:8000])  # 限制长度
            
            llm_result = await self.llm.chat(
                prompt=prompt,
                model=config.DEFAULT_MODEL,
                temperature=0.1,
                max_tokens=4096
            )
            
            if not llm_result.success:
                return {
                    "success": False,
                    "data": {},
                    "error": f"LLM 调用失败: {llm_result.error}",
                    "task_id": task_id
                }
            
            # 3. 解析 JSON
            extracted_data = self._parse_json_safe(llm_result.content)
            
            duration_ms = int((time.time() - start) * 1000)
            logger.info(
                f"[EXTRACT:{task_id}] 提取完成 | 字段数={len(extracted_data)} | "
                f"耗时={duration_ms}ms | tokens={llm_result.token_usage}"
            )
            
            return {
                "success": True,
                "data": extracted_data,
                "error": None,
                "task_id": task_id,
                "duration_ms": duration_ms,
                "token_usage": llm_result.token_usage
            }
            
        except Exception as e:
            logger.error(f"[EXTRACT:{task_id}] 异常: {str(e)}", exc_info=True)
            return {
                "success": False,
                "data": {},
                "error": str(e),
                "task_id": task_id
            }
    
    # =====================================================================
    # Step 2: 合并多个参考文档的提取结果
    # =====================================================================
    
    async def merge_extractions(self, extraction_results: List[Dict]) -> Dict:
        """
        合并多个参考文档的提取结果，去重、补全
        
        策略:
          - 相同字段：优先用置信度高的（或字符数多的）
          - 列表字段：合并去重
        """
        merged = {}
        
        for result in extraction_results:
            if not result.get("success"):
                continue
            data = result.get("data", {})
            
            for key, value in data.items():
                if key not in merged:
                    merged[key] = value
                elif isinstance(value, list) and isinstance(merged[key], list):
                    # 列表合并去重
                    merged[key] = list(dict.fromkeys(merged[key] + value))
                elif isinstance(value, str) and len(str(value)) > len(str(merged[key])):
                    # 字符串优先用更长的（通常更完整）
                    merged[key] = value
        
        logger.info(f"[MERGE] 合并完成 | 来源文档={len(extraction_results)} | 总字段数={len(merged)}")
        return merged
    
    # =====================================================================
    # Step 3: 匹配表格字段并生成填写值 (替代 Dify form-generate)
    # =====================================================================
    
    async def generate_fill_values(
        self,
        extracted_data: Dict,
        template_path: str,
        task_id: str
    ) -> Dict:
        """
        根据提取信息和空白表格模板，生成字段填写值
        
        Args:
            extracted_data: 已提取的键值对
            template_path: 空白表格模板路径 (.docx / .xlsx)
            task_id: 任务ID
            
        Returns:
            {"success": bool, "fill_values": dict, "error": str}
        """
        logger.info(f"[GENERATE:{task_id}] 开始生成填写值")
        start = time.time()
        
        try:
            # 1. 从模板提取字段列表
            field_list = self._extract_fields_from_template(template_path)
            logger.info(f"[GENERATE:{task_id}] 模板字段: {field_list}")
            
            if not field_list:
                return {
                    "success": False,
                    "fill_values": {},
                    "error": "无法从模板提取字段列表"
                }
            
            # 2. 调用 LLM 匹配字段
            prompt = FORM_GENERATE_PROMPT.format(
                extracted_json=json.dumps(extracted_data, ensure_ascii=False, indent=2),
                field_list=json.dumps(field_list, ensure_ascii=False)
            )
            
            llm_result = await self.llm.chat(
                prompt=prompt,
                model=config.DEFAULT_MODEL,
                temperature=0.1,
                max_tokens=4096
            )
            
            if not llm_result.success:
                return {
                    "success": False,
                    "fill_values": {},
                    "error": f"LLM 调用失败: {llm_result.error}"
                }
            
            # 3. 解析填写值
            fill_values = self._parse_json_safe(llm_result.content)
            
            # 4. 后处理：确保所有模板字段都有值（找不到的填"[待补充]"）
            for field in field_list:
                if field not in fill_values:
                    fill_values[field] = "[待补充]"
            
            duration_ms = int((time.time() - start) * 1000)
            filled_count = sum(1 for v in fill_values.values() if v != "[待补充]")
            
            logger.info(
                f"[GENERATE:{task_id}] 生成完成 | 总字段={len(field_list)} | "
                f"已填={filled_count} | 待补充={len(field_list)-filled_count} | "
                f"耗时={duration_ms}ms"
            )
            
            return {
                "success": True,
                "fill_values": fill_values,
                "error": None,
                "duration_ms": duration_ms,
                "token_usage": llm_result.token_usage,
                "fill_rate": round(filled_count / len(field_list) * 100, 1) if field_list else 0
            }
            
        except Exception as e:
            logger.error(f"[GENERATE:{task_id}] 异常: {str(e)}", exc_info=True)
            return {
                "success": False,
                "fill_values": {},
                "error": str(e)
            }
    
    # =====================================================================
    # Step 4: 渲染表格 (调用现有 renderer)
    # =====================================================================
    
    def render_document(
        self,
        template_path: str,
        output_path: str,
        fill_values: Dict
    ) -> Dict:
        """
        调用 docx/xlsx renderer 生成填写后的文件
        
        Returns:
            {"success": bool, "output_path": str, "filled_count": int, "error": str}
        """
        ext = os.path.splitext(template_path)[1].lower()
        
        # 预处理：将所有值转换为字符串（Excel 不支持列表直接写入）
        processed_values = {}
        for key, value in fill_values.items():
            if isinstance(value, list):
                # 列表转为逗号分隔字符串
                processed_values[key] = ", ".join(str(v) for v in value)
            elif value is None:
                processed_values[key] = ""
            else:
                processed_values[key] = str(value)
        
        try:
            if ext == ".docx":
                from app.services.docx_renderer import fill_word_document
                result = fill_word_document(template_path, output_path, processed_values)
            elif ext in [".xlsx", ".xls"]:
                from app.services.xlsx_renderer import fill_excel_document
                result = fill_excel_document(template_path, output_path, processed_values)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件类型: {ext}",
                    "output_path": None
                }
            
            logger.info(
                f"[RENDER] 渲染完成 | {ext} | 填充={result.get('filledCount', 0)} | "
                f"输出={output_path}"
            )
            return result
            
        except Exception as e:
            logger.error(f"[RENDER] 渲染异常: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "output_path": None
            }
    
    # =====================================================================
    # 端到端：一次调用完成提取+生成+渲染
    # =====================================================================
    
    async def fill_form_end_to_end(
        self,
        reference_paths: List[str],
        template_path: str,
        output_path: str,
        task_id: str
    ) -> Dict:
        """
        端到端智能填表
        
        流程:
          1. 从所有参考文档提取信息
          2. 合并提取结果
          3. 匹配模板字段生成填写值
          4. 渲染输出文件
        
        Returns:
            {
                "success": bool,
                "output_path": str,
                "fill_values": dict,
                "fill_rate": float,
                "error": str
            }
        """
        logger.info(f"[E2E:{task_id}] 开始端到端填表 | 参考文档={len(reference_paths)} | 模板={template_path}")
        
        # Step 1: 提取
        extractions = []
        for path in reference_paths:
            result = await self.extract_from_document(path, f"{task_id}_ref")
            extractions.append(result)
        
        # Step 2: 合并
        merged = await self.merge_extractions(extractions)
        if not merged:
            return {
                "success": False,
                "output_path": None,
                "fill_values": {},
                "fill_rate": 0,
                "error": "未能从参考文档提取任何信息"
            }
        
        # Step 3: 生成填写值
        gen_result = await self.generate_fill_values(merged, template_path, task_id)
        if not gen_result.get("success"):
            return {
                "success": False,
                "output_path": None,
                "fill_values": gen_result.get("fill_values", {}),
                "fill_rate": 0,
                "error": gen_result.get("error", "生成填写值失败")
            }
        
        # Step 4: 渲染
        render_result = self.render_document(
            template_path,
            output_path,
            gen_result["fill_values"]
        )
        
        return {
            "success": render_result.get("success", False),
            "output_path": render_result.get("outputPath") or render_result.get("output_path"),
            "fill_values": gen_result["fill_values"],
            "fill_rate": gen_result.get("fill_rate", 0),
            "error": render_result.get("error")
        }
    
    # =====================================================================
    # 内部工具方法
    # =====================================================================
    
    def _extract_text(self, file_path: str) -> str:
        """从文件中提取纯文本（非 OCR）"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".docx":
            return self._extract_docx_text(file_path)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == ".pdf":
            return self._extract_pdf_text(file_path)
        else:
            logger.warning(f"[EXTRACT] 不支持的文件类型: {ext}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """用 python-docx 提取 Word 文本"""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 也提取表格文本
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_texts.append(cell.text.strip())
            
            all_text = "\n".join(paragraphs + table_texts)
            return all_text
        except Exception as e:
            logger.error(f"[EXTRACT] docx 提取失败: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """用 PyPDF2 / pdfplumber 提取 PDF 文本"""
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                texts = [page.extract_text() or "" for page in pdf.pages]
                return "\n".join(texts)
        except ImportError:
            logger.warning("[EXTRACT] pdfplumber 未安装，尝试 PyPDF2")
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    texts = [page.extract_text() or "" for page in reader.pages]
                    return "\n".join(texts)
            except Exception as e2:
                logger.error(f"[EXTRACT] PDF 提取失败: {e2}")
                return ""
        except Exception as e:
            logger.error(f"[EXTRACT] PDF 提取失败: {e}")
            return ""
    
    def _extract_fields_from_template(self, template_path: str) -> List[str]:
        """从空白表格模板提取字段名列表"""
        ext = os.path.splitext(template_path)[1].lower()
        fields = []
        
        if ext == ".docx":
            try:
                import docx
                doc = docx.Document(template_path)
                
                # 从表格提取：左列通常是字段名
                for table in doc.tables:
                    for row in table.rows:
                        for i, cell in enumerate(row.cells):
                            text = cell.text.strip()
                            # 假设第一列是字段名，第二列是值
                            if i == 0 and text and not text.startswith("["):
                                fields.append(text)
                
                # 去重
                fields = list(dict.fromkeys(fields))
            except Exception as e:
                logger.error(f"[FIELD] docx 字段提取失败: {e}")
        
        elif ext in [".xlsx", ".xls"]:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(template_path)
                ws = wb.active
                
                # 从第一行/第一列提取字段名
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value:
                            text = str(cell.value).strip()
                            if text and not text.startswith("[") and "请" not in text:
                                fields.append(text)
                
                fields = list(dict.fromkeys(fields))
            except Exception as e:
                logger.error(f"[FIELD] xlsx 字段提取失败: {e}")
        
        logger.info(f"[FIELD] 提取字段: {fields}")
        return fields
    
    def _parse_json_safe(self, text: str) -> Dict:
        """安全解析 LLM 返回的 JSON，处理各种格式问题"""
        if not text:
            return {}
        
        text = text.strip()
        
        # 移除 markdown 代码块
        if text.startswith("```"):
            lines = text.split("\n")
            if lines[0].strip().startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            text = "\n".join(lines)
        
        # 尝试提取 JSON 对象
        try:
            return json.loads(text.strip())
        except json.JSONDecodeError:
            # 尝试找最外层的大括号
            match = re.search(r'\{.*\}', text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except json.JSONDecodeError:
                    pass
            
            logger.error(f"[PARSE] JSON 解析失败，原始文本前200字: {text[:200]}")
            return {}
