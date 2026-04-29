# -*- coding: utf-8 -*-
"""
=============================================================================
LingDoc AI Service - OCR Engine Service
=============================================================================
用途：封装 PaddleOCR 引擎，提供文档（PDF/Word/图片）到文本的转换能力

设计原则：
  1. 单例模式：PaddleOCR 引擎初始化耗时，全局只创建一次
  2. 多格式支持：PDF/Word 自动转图片后识别
  3. 结构化输出：返回文本、位置、置信度、页码等完整信息
  4. 异步安全：PaddleOCR 非线程安全，使用锁保护

性能说明：
  - 首次初始化：约 3-5 秒（下载/加载模型）
  - 单页识别：约 0.5-2 秒（取决于图片大小和文字密度）
  - CPU 占用：单核满载，建议 Docker 限制 CPU 或开启多进程

对接说明（给主项目工程师）：
  - 主后端不需要关心 OCR 细节，只需传文件路径
  - AI 服务内部完成：PDF 拆页 → 转图片 → OCR 识别 → 合并结果
  - 大文件（>50页）建议主后端设置超时 ≥ 120 秒
=============================================================================
"""

import os
import io
import time
import asyncio
import tempfile
import hashlib
from pathlib import Path
from typing import Dict, List, Optional, Any
from functools import lru_cache

# PaddleOCR 依赖（延迟导入，避免服务启动时卡住）
try:
    from paddleocr import PaddleOCR
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    PaddleOCR = None

# 文档处理
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    fitz = None

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Document = None

try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False
    Image = None

from app.config import config
from app.utils.logger import logger


class OcrEngine:
    """
    PaddleOCR 引擎封装类（单例）
    
    用途：提供统一的文档 OCR 接口，隐藏底层复杂度
    
    使用方式：
        ```python
        engine = OcrEngine()
        result = await engine.process("/path/to/file.pdf")
        print(result["ocr_text"])  # 完整文本
        print(result["page_count"])  # 页数
        ```
    
    线程安全：
      - __init__ 中创建引擎实例（单线程）
      - process 方法使用 asyncio.Lock 保护（多请求并发时串行执行）
    """
    
    _instance: Optional["OcrEngine"] = None
    _lock = asyncio.Lock()
    
    def __new__(cls) -> "OcrEngine":
        """单例模式：确保全局只有一个引擎实例"""
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        """初始化 OCR 引擎（只执行一次）"""
        if self._initialized:
            return
        
        if not PADDLEOCR_AVAILABLE:
            logger.warning("PaddleOCR 未安装，OCR 功能不可用。请运行: pip install paddleocr paddlepaddle")
            self._ocr = None
            self._initialized = True
            return
        
        logger.info("正在初始化 PaddleOCR 引擎（首次加载约 3-5 秒）...")
        start = time.time()
        
        # PaddleOCR 配置
        # lang="ch"：中文+英文混合识别
        # use_angle_cls=True：启用方向分类（支持旋转图片）
        # use_gpu=False：CPU 运行（服务器通常无 GPU）
        self._ocr = PaddleOCR(
            lang=config.OCR_LANGUAGE,
            use_angle_cls=True,
            use_gpu=False,
            show_log=False,  # 禁用 Paddle 内部日志，避免污染
        )
        
        self._initialized = True
        logger.info(f"PaddleOCR 引擎初始化完成，耗时 {time.time()-start:.2f}s")
    
    async def process(self, file_path: str, task_id: str = "") -> Dict[str, Any]:
        """
        处理单个文件，返回 OCR 结果
        
        参数：
            file_path: 文件绝对路径（PDF/Word/图片）
            task_id: 任务 ID，用于日志追踪
        
        返回：
            Dict: {
                "success": bool,
                "ocr_text": str,        # 完整文本
                "page_count": int,      # 页数
                "total_lines": int,     # 文本行数
                "total_chars": int,     # 字符数
                "pages": List[Dict],    # 每页详细结果
                "process_time_ms": int,
                "error": Optional[str]
            }
        
        异常：
            不会抛出异常，错误信息封装在返回值的 error 字段中
        """
        start_time = time.time()
        
        async with OcrEngine._lock:
            logger.info(f"[OCR:{task_id}] 开始处理: {file_path}")
            
            if not PADDLEOCR_AVAILABLE:
                return {
                    "success": False,
                    "ocr_text": "",
                    "error": "PaddleOCR 未安装，无法执行 OCR 识别。请运行: pip install paddleocr paddlepaddle",
                    "process_time_ms": 0
                }
            
            try:
                # 检查文件存在性
                if not os.path.exists(file_path):
                    return {
                        "success": False,
                        "ocr_text": "",
                        "error": f"文件不存在: {file_path}",
                        "process_time_ms": 0
                    }
                
                # 检查文件大小（防止内存溢出）
                file_size = os.path.getsize(file_path)
                if file_size > 200 * 1024 * 1024:  # 200MB 上限
                    return {
                        "success": False,
                        "ocr_text": "",
                        "error": f"文件过大 ({file_size/1024/1024:.1f}MB > 200MB 限制)",
                        "process_time_ms": 0
                    }
                
                # 根据文件类型选择处理方式
                ext = Path(file_path).suffix.lower()
                
                if ext in [".pdf"]:
                    result = await self._process_pdf(file_path, task_id)
                elif ext in [".docx", ".doc"]:
                    result = await self._process_word(file_path, task_id)
                elif ext in [".png", ".jpg", ".jpeg", ".bmp"]:
                    result = await self._process_image(file_path, task_id)
                elif ext in [".txt", ".md", ".csv"]:
                    # 纯文本文件直接读取，无需 OCR
                    result = await self._process_text(file_path, task_id)
                else:
                    return {
                        "success": False,
                        "ocr_text": "",
                        "error": f"不支持的文件格式: {ext}",
                        "process_time_ms": 0
                    }
                
                # 添加处理时间
                result["process_time_ms"] = int((time.time() - start_time) * 1000)
                
                logger.info(
                    f"[OCR:{task_id}] 完成 | 页数={result.get('page_count', 0)} | "
                    f"字符={result.get('total_chars', 0)} | "
                    f"耗时={result['process_time_ms']}ms"
                )
                
                return result
                
            except Exception as e:
                logger.error(f"[OCR:{task_id}] 处理异常: {str(e)}", exc_info=True)
                return {
                    "success": False,
                    "ocr_text": "",
                    "error": f"OCR 引擎内部错误: {str(e)}",
                    "process_time_ms": int((time.time() - start_time) * 1000)
                }
    
    # -------------------------------------------------------------------------
    # 私有方法：按文件类型处理
    # -------------------------------------------------------------------------
    
    async def _process_pdf(self, file_path: str, task_id: str) -> Dict[str, Any]:
        """处理 PDF 文件：逐页转图片 → OCR 识别"""
        pages = []
        all_text = []
        
        # 使用 PyMuPDF 打开 PDF
        doc = fitz.open(file_path)
        total_pages = len(doc)
        
        for page_idx in range(total_pages):
            page = doc[page_idx]
            
            # PDF 页转图片（DPI 决定清晰度）
            pix = page.get_pixmap(matrix=fitz.Matrix(config.PDF_DPI/72, config.PDF_DPI/72))
            
            # 转 PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # 临时保存图片（PaddleOCR 需要文件路径）
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                img.save(tmp.name, "PNG")
                tmp_path = tmp.name
            
            try:
                # OCR 识别
                ocr_result = self._ocr.ocr(tmp_path, cls=True)
                
                # 解析结果
                page_text = []
                page_lines = 0
                for line in ocr_result[0] or []:
                    if line:
                        text = line[1][0]  # 文本内容
                        confidence = line[1][1]  # 置信度
                        
                        # 过滤低置信度结果
                        if confidence >= config.OCR_THRESHOLD:
                            page_text.append(text)
                            page_lines += 1
                
                page_content = "\n".join(page_text)
                all_text.append(page_content)
                
                pages.append({
                    "page_no": page_idx + 1,
                    "text": page_content,
                    "lines": page_lines,
                    "chars": len(page_content),
                })
                
            finally:
                # 清理临时文件
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
        
        doc.close()
        
        full_text = "\n\n".join(all_text)
        
        return {
            "success": True,
            "ocr_text": full_text,
            "page_count": total_pages,
            "total_lines": sum(p["lines"] for p in pages),
            "total_chars": len(full_text),
            "pages": pages,
        }
    
    async def _process_word(self, file_path: str, task_id: str) -> Dict[str, Any]:
        """处理 Word 文件：逐段提取文本（优先直接读，失败则转图片 OCR）"""
        try:
            # 方式 1：直接提取文本（速度快，格式准确）
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # 尝试提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            full_text = "\n".join(paragraphs)
            
            return {
                "success": True,
                "ocr_text": full_text,
                "page_count": 1,  # Word 页数需要额外计算，暂定为 1
                "total_lines": len(paragraphs),
                "total_chars": len(full_text),
                "pages": [{
                    "page_no": 1,
                    "text": full_text,
                    "lines": len(paragraphs),
                    "chars": len(full_text),
                }],
                "extract_method": "direct",  # 直接提取，非 OCR
            }
            
        except Exception as e:
            # 方式 2：直接提取失败，回退到 OCR（Word 转 PDF 再 OCR）
            logger.warning(f"[OCR:{task_id}] Word 直接提取失败，回退到 OCR: {str(e)}")
            # TODO: 实现 Word → PDF → OCR 回退逻辑
            return {
                "success": False,
                "ocr_text": "",
                "error": f"Word 处理失败: {str(e)}"
            }
    
    async def _process_image(self, file_path: str, task_id: str) -> Dict[str, Any]:
        """处理图片文件：直接 OCR 识别"""
        if not PADDLEOCR_AVAILABLE or self._ocr is None:
            return {
                "success": False,
                "ocr_text": "",
                "error": "PaddleOCR 未安装，无法识别图片",
            }
        
        ocr_result = self._ocr.ocr(file_path, cls=True)
        
        page_text = []
        page_lines = 0
        for line in ocr_result[0] or []:
            if line:
                text = line[1][0]
                confidence = line[1][1]
                
                if confidence >= config.OCR_THRESHOLD:
                    page_text.append(text)
                    page_lines += 1
        
        full_text = "\n".join(page_text)
        
        return {
            "success": True,
            "ocr_text": full_text,
            "page_count": 1,
            "total_lines": page_lines,
            "total_chars": len(full_text),
            "pages": [{
                "page_no": 1,
                "text": full_text,
                "lines": page_lines,
                "chars": len(full_text),
            }],
        }
    
    async def _process_text(self, file_path: str, task_id: str) -> Dict[str, Any]:
        """处理纯文本文件：直接读取内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            lines = content.split('\n')
            
            return {
                "success": True,
                "ocr_text": content,
                "page_count": 1,
                "total_lines": len(lines),
                "total_chars": len(content),
                "pages": [{
                    "page_no": 1,
                    "text": content,
                    "lines": len(lines),
                    "chars": len(content),
                }],
                "extract_method": "direct_read",
            }
        except Exception as e:
            return {
                "success": False,
                "ocr_text": "",
                "error": f"文本读取失败: {str(e)}",
            }


# =============================================================================
# 快捷函数（给路由层直接调用）
# =============================================================================

async def ocr_process(file_path: str, task_id: str = "") -> Dict[str, Any]:
    """
    快捷函数：OCR 处理入口
    
    用途：路由层直接调用，无需关心引擎初始化细节
    
    示例：
        ```python
        result = await ocr_process("/uploads/report.pdf", task_id="t_001")
        if result["success"]:
            print(result["ocr_text"])
        ```
    """
    engine = OcrEngine()
    return await engine.process(file_path, task_id)
